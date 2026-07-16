"""
HW2-specific Markdown -> docx converter.

Unlike HW1's md2docx.py (which uses Word's built-in Heading styles -- a
different font and size per level -- and Consolas for code), this renderer
enforces exactly one font family and two font sizes total: one for headings
(all levels alike) and one shared by everything else (paragraphs, list
items, table cells, and inline/block code).
"""
import sys, re, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src, dst = sys.argv[1], sys.argv[2]
srcdir = os.path.dirname(os.path.abspath(src))
with open(src, encoding="utf-8") as f:
    lines = f.read().split("\n")

FONT_NAME = "Arial"
TITLE_SIZE = Pt(20)
HEADER_SIZE = Pt(14)
BODY_SIZE = Pt(11)

doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = BODY_SIZE

INLINE = re.compile(r"(`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_runs(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            r = p.add_run(m.group(1) if m else tok)
        else:
            r = p.add_run(tok)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE


def add_heading(text, is_title=False):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = TITLE_SIZE if is_title else HEADER_SIZE
    r.bold = True


def add_code_block(block_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for i, ln in enumerate(block_lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE


def add_table(rows):
    cells = [[c.strip() for c in row.strip().strip("|").split("|")] for row in rows]
    header, body_rows = cells[0], cells[2:]  # cells[1] is the |---| separator
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], h)
    for row in body_rows:
        cs = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(cs):
                cs[j].paragraphs[0].text = ""
                add_runs(cs[j].paragraphs[0], val)


i = 0
n = len(lines)
while i < n:
    line = lines[i]
    if line.startswith("#"):
        add_heading(re.sub(r"^#+\s*", "", line).strip(), is_title=line.startswith("# "))
    elif line.strip().startswith("```"):
        block = []
        i += 1
        while i < n and not lines[i].strip().startswith("```"):
            block.append(lines[i])
            i += 1
        add_code_block(block)
    elif line.strip().startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append(lines[i])
            i += 1
        add_table(rows)
        continue
    elif line.strip().startswith("!["):
        m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
        img = os.path.normpath(os.path.join(srcdir, m.group(1)))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img, width=Inches(5.6))
    elif line.strip().startswith("- "):
        p = doc.add_paragraph()
        add_runs(p, line.strip()[2:])
    elif line.strip() == "":
        pass
    else:
        buf = [line.strip()]
        i += 1
        while i < n and lines[i].strip() and not re.match(r"^(#|\||```|- )", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))
        continue
    i += 1

doc.save(dst)
print("wrote", dst)
