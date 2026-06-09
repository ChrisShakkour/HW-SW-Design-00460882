import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src, dst = sys.argv[1], sys.argv[2]
with open(src, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")

def add_runs(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            r = p.add_run(m.group(1)); r.bold = True   # show link text only
        else:
            p.add_run(tok)

def add_code(block):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    for i, ln in enumerate(block):
        if i:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def add_table(rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]   # cells[1] is the |---| separator
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cs = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(cs):
                cs[j].paragraphs[0].text = ""
                add_runs(cs[j].paragraphs[0], val)

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=0)
    elif line.startswith("## "):
        doc.add_heading(re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line[3:].strip()), level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.strip().startswith("```"):
        block = []; i += 1
        while i < n and not lines[i].strip().startswith("```"):
            block.append(lines[i]); i += 1
        add_code(block)
    elif line.strip().startswith("|") and i + 1 < n and set(lines[i+1].strip()) <= set("|-: "):
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append(lines[i]); i += 1
        add_table(rows); continue
    elif line.strip().startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, line.strip()[2:])
    elif line.strip() == "---":
        pass
    elif line.strip() == "":
        pass
    else:
        # gather a wrapped paragraph (consecutive plain lines)
        buf = [line.strip()]; i += 1
        while i < n and lines[i].strip() and not re.match(r"^(#|\||```|- |---)", lines[i].strip()):
            buf.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_runs(p, " ".join(buf)); continue
    i += 1

doc.save(dst)
print("wrote", dst)
